import streamlit as st
import numpy as np
import re
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import MinMaxScaler

from utils.common_functions import (
    load_docx_from_folder,
    spacy_tokenizer,
    progress_bar_update,
    nlp
)
from utils.explanation import generate_explanation_with_llm_job_to_cv
from src.db.queries import (
    get_job_id_by_filename,
    get_job_industry,
    get_cv_industry_scores
)

cv_folder = './DataSet/cv'
job_folder = './DataSet/job_descriptions'

def get_keyword_matching_scores(custom_skills, cv_text): ##
    skill_matches = {}
    for skill, weight in custom_skills:
        count = len(re.findall(rf"\b{re.escape(skill.lower())}\b", cv_text.lower()))
        skill_matches[skill] = (count > 0, weight)
    total_score = sum(weight for matched, weight in skill_matches.values() if matched)
    return total_score / 100

def get_cv_keyword_matching_scores(custom_skills, cv_texts):
    scores = [get_keyword_matching_scores(custom_skills, cv_text) for cv_text in cv_texts]
    return np.array(scores)

def get_skills_matched_for_cv(custom_skills, cv_text):
    skill_matches = {}
    for skill, weight in custom_skills:
        count = len(re.findall(rf"\b{re.escape(skill.lower())}\b", cv_text.lower()))
        skill_matches[skill] = (count > 0, weight)
    return skill_matches

def get_matching_scores_between_cvs_and_job_description(cv_texts, job_text, progress_bar, status_text): ##
    if not cv_texts:
        return []

    def compute_batch_similarity(batch, job_vector):
        return [job_vector.similarity(nlp(cv_text)) for cv_text in batch]

    def chunk_list(lst, n):
        return [lst[i:i+n] for i in range(0, len(lst), n)]

    cv_texts_preprocessed = [" ".join(spacy_tokenizer(cv_text)) for cv_text in cv_texts]
    job_text_preprocessed = " ".join(spacy_tokenizer(job_text))

    progress_bar_update(30, progress_bar, status_text)

    job_vector = nlp(job_text_preprocessed)
    batch_size = 50
    chunks = chunk_list(cv_texts_preprocessed, batch_size)
    num_threads = min(8, len(chunks))

    with ThreadPoolExecutor(max_workers=num_threads) as executor:
        results = executor.map(lambda batch: compute_batch_similarity(batch, job_vector), chunks)

    similarities_embeddings = [sim for batch in results for sim in batch]

    progress_bar_update(70, progress_bar, status_text)

    combined_texts = cv_texts_preprocessed + [job_text_preprocessed]
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(combined_texts)
    cv_vectors = tfidf_matrix[:len(cv_texts_preprocessed)]
    job_vectors = tfidf_matrix[len(cv_texts_preprocessed):]

    similarities_tfidf = cosine_similarity(job_vectors, cv_vectors)
    progress_bar_update(80, progress_bar, status_text)

    scaler = MinMaxScaler()
    similarities_tfidf_scaled = scaler.fit_transform(np.array(similarities_tfidf).reshape(-1, 1)).flatten()

    progress_bar_update(90, progress_bar, status_text)
    similarities_embeddings = np.array(similarities_embeddings)

    final_embeddings_tfidf_scores = 0.6 * similarities_embeddings + 0.4 * similarities_tfidf_scaled
    return final_embeddings_tfidf_scores

def filter_cvs_by_industry(db_path, cv_folder, selected_industry):
    cv_texts, cv_filenames, _ = load_docx_from_folder(cv_folder, is_cv=True)
    
    industry_scores = get_cv_industry_scores(db_path, cv_filenames, selected_industry)
    
    filtered_indices = [i for i, score in enumerate(industry_scores) if score > 0]
    
    if not filtered_indices:
        return [], [], np.array([])
        
    return (
        [cv_texts[i] for i in filtered_indices],
        [cv_filenames[i] for i in filtered_indices],
        industry_scores[filtered_indices]
    )

st.title("📌 Match Job → CVs")

st.subheader("🛠️ Define Required Skills and Weights")
skill_count = st.number_input("How many skills do you want to assess?", min_value=1, max_value=20, value=3)

custom_skills = []
total_weight = 0

if skill_count:
    with st.form("skill_form"):
        for i in range(skill_count):
            col1, col2 = st.columns([2, 1])
            with col1:
                skill_name = st.text_input(f"Skill #{i + 1} name", key=f"skill_name_{i}")
            with col2:
                weight = st.number_input(f"Weight % for Skill #{i + 1}", min_value=0, max_value=100,
                                         key=f"skill_weight_{i}")

            if skill_name:
                custom_skills.append((skill_name.strip(), weight))
                total_weight += weight

        if total_weight != 100:
            st.error(f"Total weights must sum to 100% (current sum: {total_weight}%)")
        else:
            st.success("Weights sum to 100% - Ready to proceed!")

        submitted = st.form_submit_button("✅ Done selecting skills")

uploaded_file = st.file_uploader("Upload a Job Description", type=["docx"])
progress_bar = st.progress(0)
status_text = st.empty()

col1, col2, col3 = st.columns([1,2,1])
with col2:
    start_button = st.button(
        "Get Matching CVs",
        type="primary",
        use_container_width=True
    )

if start_button:
    if not uploaded_file:
        st.error("❌ Please upload a job description file")
    elif total_weight != 100:
        st.error(f"❌ Total weight must be exactly 100%. Current sum: {total_weight}%.")
    else:
        job_filename = uploaded_file.name
        job_id = get_job_id_by_filename("./data/cvs_metadata.sqlite", job_filename)

        if job_id is None:
            st.error("❌ Job description not found in database.")
        else:
            selected_industry = get_job_industry("./data/cvs_metadata.sqlite", job_id)
            if selected_industry is None:
                st.error("❌ Industry for the uploaded job not found in database.")
            else:
                st.success(f"Industry detected: {selected_industry}")

                doc = Document(uploaded_file)
                text = ' '.join([para.text for para in doc.paragraphs])
                job_text = text.replace('\n', ' ').replace('  ', ' ').split("Benefits:")[0]

                # Load CVs
                cv_texts, cv_filenames, industry_scores = filter_cvs_by_industry(
                    "./data/cvs_metadata.sqlite",
                    cv_folder,
                    selected_industry
                )

                progress_bar_update(10, progress_bar, status_text)

                skills_scores = get_cv_keyword_matching_scores(custom_skills, cv_texts)
                semantic_scores = get_matching_scores_between_cvs_and_job_description(cv_texts, job_text, progress_bar, status_text)

                final_scores = (
                    0.1 * industry_scores +
                    0.3 * skills_scores +
                    0.6 * semantic_scores
                )

                st.markdown("---")
                st.subheader("🎯 Best CVs Match")
                final_scores_indices = np.argsort(final_scores)[::-1][:5]
                for i in final_scores_indices:
                    st.write(f"{cv_filenames[i]} - Matching score: {final_scores[i] * 100:.2f}%")

                doc = Document(cv_folder + '/' + cv_filenames[final_scores_indices[0]])
                best_cv = '\n'.join([para.text for para in doc.paragraphs])

                st.markdown("---")
                st.header("The top matching CV is:")
                st.write(best_cv)

                st.markdown("---")
                st.subheader("Matched Skills:")
                common_skills = [
                    skill for skill, _ in custom_skills
                    if re.search(rf'\b{re.escape(skill.lower())}\b', best_cv.lower())
                ]
                if common_skills:
                    st.write(", ".join(common_skills))
                else:
                    st.write("No matched skills found.")

                st.markdown("---")
                st.header("Selection Explanation:")
                selected_cv_index = final_scores_indices[0]
                st.write(
                    f"Industry knowledge score: {industry_scores[selected_cv_index] * 100:.2f}% | "
                    f"Technical Qualification score: {skills_scores[selected_cv_index] * 100:.2f}% | "
                    f"Job - CV Matching score: {semantic_scores[selected_cv_index] * 100:.2f}%"
                )

                skills_matched = get_skills_matched_for_cv(custom_skills, cv_texts[selected_cv_index])

                explanation_text = generate_explanation_with_llm_job_to_cv(
                    cv_filename=cv_filenames[selected_cv_index],
                    domain_score=industry_scores[selected_cv_index],
                    skills_score=skills_scores[selected_cv_index],
                    matching_score=semantic_scores[selected_cv_index],
                    matched_skills=skills_matched,
                    domain_selected=selected_industry
                )
                st.subheader("Explanation for the Best Match")
                st.write(explanation_text)

                progress_bar_update(100, progress_bar, status_text)
                st.balloons()